from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
import shutil


class DocxGenerator:
    def __init__(self):
        self.document = Document()
        self.styles = self.document.styles  # Define styles for potential future customization

    def generate(self, report_content: str, filename: str = "progress_report.docx"):
        lines = report_content.split('\n')
        processed_lines = []

        for line in lines:
            # Check for bold text (assuming single leading/trailing asterisks or underscores)
            if (line.strip().startswith("*") and line.strip().endswith("*")) or \
               (line.strip().startswith("_") and line.strip().endswith("_")):
                text = line.strip()[1:-1]  # Extract text between symbols
                processed_lines.append(text)  # Add text for bold formatting
            # Check for italics (assuming single leading/trailing asterisks or underscores, excluding bold)
            elif (line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**")) or \
                 (line.strip().startswith("_") and line.strip().endswith("_") and not line.strip().startswith("__")):
                text = line.strip()[1:-1]  # Extract text between symbols
                processed_lines.append(f"*{text}*")  # Add italic formatting
            # Check for headings (assuming single leading "#")
            elif line.strip().startswith("#"):
                text = line.strip()[1:]  # Extract text after #
                processed_lines.append(text)  # Add text for heading formatting
            # Check for lists (identify by starting hyphen, plus, or number)
            elif line.strip().startswith("-") or line.strip().startswith("+") or line.strip().isdigit():
                processed_lines.append(line.strip())  # Keep list formatting
            else:
                processed_lines.append(line.strip())  # No special formatting, keep as is

        # Process the formatted lines
        for line in processed_lines:
            # Apply formatting based on processed text
            if line.startswith("*"):
                self.add_bold_text(line)
            elif line.strip().startswith("#"):
                self.add_heading(line)
            else:
                self.add_normal_text(line)

        # Save the document
        try:
            current_dir = os.path.abspath(os.getcwd())
            file_path = os.path.join(current_dir, filename)

            # First, try to save to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                self.document.save(tmp_file.name)

            # If temporary save succeeds, try to copy the file to the desired location
            try:
                shutil.copy2(tmp_file.name, file_path)
                os.unlink(tmp_file.name)  # Remove the temporary file
                print(f"Report saved successfully as: {file_path}")
            except PermissionError:
                print(f"Permission denied when saving to {file_path}")
                print(f"The report has been saved as a temporary file: {tmp_file.name}")
                print("Please close any programs that might be using the target file and try again.")
            except Exception as e:
                print(f"Error copying the file: {str(e)}")
                print(f"The report has been saved as a temporary file: {tmp_file.name}")
        except Exception as e:
            print(f"Error generating the document: {str(e)}")

    def add_heading(self, text: str):
        """Adds a bold heading with increased font size and centered alignment."""
        p = self.document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_bold_text(self, text: str):
        """Adds bold text with some spacing."""
        p = self.document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.paragraph_format.space_before = Pt(6)  # Add some space before bold text

    def add_normal_text(self, text: str):
        """Adds normal text with some spacing after previous paragraph and sets font size."""
        p = self.document.add_paragraph()
        run = p.add_run(text)
        p.paragraph_format.space_before = Pt(6)  # Add some space before normal text
        run.font.size = Pt(12)